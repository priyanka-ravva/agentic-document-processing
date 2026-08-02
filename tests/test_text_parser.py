import json
import os
import tempfile

import docx
import openpyxl

from src.tools.text_parser import extract_text_from_text_document


def test_extracts_text_from_txt_file() -> None:
    with tempfile.NamedTemporaryFile(suffix='.txt', mode='w+', encoding='utf-8', delete=False) as tmp:
        tmp.write('Hello world\nThis is a test.')
        tmp_path = tmp.name

    try:
        extracted = extract_text_from_text_document(tmp_path)
        assert 'Hello world' in extracted
        assert 'This is a test.' in extracted
    finally:
        os.unlink(tmp_path)


def test_extracts_text_from_json_file() -> None:
    data = {'invoice': {'number': 'INV-001', 'total': 100}}
    with tempfile.NamedTemporaryFile(suffix='.json', mode='w+', encoding='utf-8', delete=False) as tmp:
        json.dump(data, tmp)
        tmp_path = tmp.name

    try:
        extracted = extract_text_from_text_document(tmp_path)
        assert 'INV-001' in extracted
        assert '100' in extracted
        assert '{' in extracted and '}' in extracted
    finally:
        os.unlink(tmp_path)


def test_extracts_text_from_csv_file() -> None:
    with tempfile.NamedTemporaryFile(suffix='.csv', mode='w+', encoding='utf-8', delete=False) as tmp:
        tmp.write('name,amount\nAlice,42\nBob,99')
        tmp_path = tmp.name

    try:
        extracted = extract_text_from_text_document(tmp_path)
        assert 'name | amount' in extracted
        assert 'Alice | 42' in extracted
        assert 'Bob | 99' in extracted
    finally:
        os.unlink(tmp_path)


def test_extracts_text_from_xlsx_file() -> None:
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = os.path.join(tmp_dir, 'sheet.xlsx')
        workbook = openpyxl.Workbook()
        sheet = workbook.active
        sheet.title = 'Sheet1'
        sheet.append(['name', 'score'])
        sheet.append(['Alice', 95])
        workbook.save(tmp_path)

        extracted = extract_text_from_text_document(tmp_path)
        assert '--- Sheet: Sheet1 ---' in extracted
        assert 'name | score' in extracted
        assert 'Alice | 95' in extracted


def test_extracts_text_from_docx_file() -> None:
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = os.path.join(tmp_dir, 'document.docx')
        document = docx.Document()
        document.add_paragraph('First paragraph.')
        document.add_paragraph('Second paragraph.')
        document.save(tmp_path)

        extracted = extract_text_from_text_document(tmp_path)
        assert 'First paragraph.' in extracted
        assert 'Second paragraph.' in extracted
